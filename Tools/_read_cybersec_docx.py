# -*- coding: utf-8 -*-
import sys
from pathlib import Path

from docx import Document

SRC = Path(
    r"C:\Users\lie76\Desktop\三维重建软件文档\过审材料\软件资料需求2026.06.23"
    r"\医学影像三维重建软件-网络安全测试用例-修订.docx"
)
OUT = Path(__file__).resolve().parent / "_cybersec_docx_dump.txt"

doc = Document(str(SRC))
lines = []
for para in doc.paragraphs:
    if para.text.strip():
        lines.append(f"P: {para.text}")
lines.append("---TABLE---")
for ti, table in enumerate(doc.tables):
    lines.append(f"TABLE {ti}: {len(table.rows)}x{len(table.columns)}")
    for ri, row in enumerate(table.rows):
        cells = [c.text.replace("\n", " | ") for c in row.cells]
        lines.append(f"R{ri}: " + " || ".join(cells))

OUT.write_text("\n".join(lines), encoding="utf-8")
print(OUT)
