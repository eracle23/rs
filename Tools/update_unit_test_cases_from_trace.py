"""
根据《软件需求规格说明书.docx》与《1.1单元测试用例_修改版.docx》中的需求追溯，
自动回写/补齐《1.1单元测试用例.docx》的“需求追溯”列与附录追溯矩阵。

输出新文件，避免覆盖原件（防止 Word/WPS 占用导致 PermissionError）。
"""

from __future__ import annotations

from dataclasses import dataclass
from datetime import date
from pathlib import Path
import re
from typing import Iterable

from docx import Document


ROOT = Path.home() / "Desktop" / "三维重建软件文档" / "过审材料"
SRC_11 = ROOT / "1.1单元测试用例.docx"
SRC_11_MOD = ROOT / "1.1单元测试用例_修改版.docx"


REQ_ID_RE = re.compile(r"\b(?:FR|DR|PR|SR|RR)-\d{3}\b")


@dataclass(frozen=True)
class Requirement:
    rid: str
    desc: str


def _iter_table_text(doc: Document) -> Iterable[str]:
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                txt = (cell.text or "").strip()
                if txt:
                    yield txt


def _find_srs_doc() -> Path:
    """
    在 ROOT 下自动挑一份“最像 SRS”的文档：
    - 能解析为 docx
    - 含最多需求编号（FR/DR/PR/SR/RR-xxx）
    """
    best: tuple[int, Path] | None = None
    for f in sorted(ROOT.glob("*.docx")):
        if f.name.startswith("~$"):
            continue
        try:
            doc = Document(str(f))
        except Exception:
            continue
        ids = set()
        for p in doc.paragraphs:
            ids.update(REQ_ID_RE.findall(p.text or ""))
        for txt in _iter_table_text(doc):
            ids.update(REQ_ID_RE.findall(txt))
        score = len(ids)
        if score == 0:
            continue
        if best is None or score > best[0]:
            best = (score, f)
    if not best:
        raise FileNotFoundError(f"未在目录中找到包含需求编号的 docx：{ROOT}")
    return best[1]


def extract_requirements_from_srs(srs_path: Path) -> dict[str, Requirement]:
    """
    从 SRS 中尽量提取“需求编号 -> 描述”。
    主要从表格行中抽取：若某行任一单元格包含 RID，则取同一行其它单元格拼成描述。
    """
    doc = Document(str(srs_path))
    reqs: dict[str, Requirement] = {}

    # 1) 表格优先（通常 SRS 的需求在表格里）
    for t in doc.tables:
        for row in t.rows:
            cells = [re.sub(r"\s+", " ", (c.text or "").strip()) for c in row.cells]
            if not any(cells):
                continue
            row_text = " | ".join(cells)
            rids = set(REQ_ID_RE.findall(row_text))
            if not rids:
                continue
            # 描述：去掉 rid 本身，取最长的非空单元格作为描述候选
            desc_candidates = [c for c in cells if c and not REQ_ID_RE.fullmatch(c)]
            desc = max(desc_candidates, key=len) if desc_candidates else row_text
            desc = REQ_ID_RE.sub("", desc).strip(" |-：:；;")
            desc = desc or "（SRS 表格行描述缺失）"
            for rid in rids:
                if rid not in reqs:
                    reqs[rid] = Requirement(rid=rid, desc=desc)

    # 2) 段落补充（少量需求会以列表/段落出现）
    for p in doc.paragraphs:
        txt = re.sub(r"\s+", " ", (p.text or "").strip())
        if not txt:
            continue
        for rid in REQ_ID_RE.findall(txt):
            if rid in reqs:
                continue
            # 粗略取 rid 后面的文字作为描述
            m = re.search(re.escape(rid) + r"\s*[-：:]\s*(.+)$", txt)
            desc = (m.group(1).strip() if m else txt.replace(rid, "").strip())
            reqs[rid] = Requirement(rid=rid, desc=desc or "（段落描述缺失）")

    return reqs


def extract_trace_from_modified_cases(mod_path: Path) -> dict[str, set[str]]:
    """
    从《1.1单元测试用例_修改版.docx》里提取 “用例编号 -> {需求编号}”。
    该文档里既有各模块表格，也有附录追溯矩阵。
    """
    doc = Document(str(mod_path))
    case_to_reqs: dict[str, set[str]] = {}

    # 表格：列名含“用例编号/需求追溯”
    for t in doc.tables:
        if len(t.rows) < 2:
            continue
        headers = [re.sub(r"\s+", "", c.text or "") for c in t.rows[0].cells]
        if not headers:
            continue
        try:
            idx_case = next(i for i, h in enumerate(headers) if "用例编号" in h)
        except StopIteration:
            continue
        idx_req = None
        for i, h in enumerate(headers):
            if "需求追溯" in h or "需求" == h or "追溯" in h:
                idx_req = i
                break
        if idx_req is None:
            continue

        for r in t.rows[1:]:
            cells = [c.text.strip() for c in r.cells]
            if idx_case >= len(cells):
                continue
            cid = re.sub(r"\s+", "", cells[idx_case])
            if not cid:
                continue
            reqs = set(REQ_ID_RE.findall(cells[idx_req] if idx_req < len(cells) else ""))
            if not reqs:
                continue
            case_to_reqs.setdefault(cid, set()).update(reqs)

    # 段落/文本兜底：抓 “UT-xxx ... FR-001” 这类同一行
    for p in doc.paragraphs:
        line = re.sub(r"\s+", " ", (p.text or "").strip())
        if not line:
            continue
        # case id 形态：UT-XX-001 / UT-001 等（尽量宽松）
        m = re.search(r"\bUT-[A-Z]{0,3}-?\d{3}\b", line)
        if not m:
            continue
        cid = m.group(0)
        reqs = set(REQ_ID_RE.findall(line))
        if reqs:
            case_to_reqs.setdefault(cid, set()).update(reqs)

    return case_to_reqs


def _set_cell_text(cell, text: str) -> None:
    # docx 单元格直接赋值会保留原 run 但可接受；这里做最简单的清空重写
    cell.text = ""
    cell.text = text


def update_original_doc(
    src_path: Path,
    case_to_reqs: dict[str, set[str]],
    reqs: dict[str, Requirement],
    out_path: Path,
) -> None:
    doc = Document(str(src_path))

    # 1) 补齐“需求追溯”列（若存在）
    for t in doc.tables:
        if len(t.rows) < 2:
            continue
        headers = [re.sub(r"\s+", "", c.text or "") for c in t.rows[0].cells]
        if not headers:
            continue
        # 识别：用例明细表
        try:
            idx_case = next(i for i, h in enumerate(headers) if "用例编号" in h)
        except StopIteration:
            continue
        idx_req = None
        for i, h in enumerate(headers):
            if "需求追溯" in h:
                idx_req = i
                break
        if idx_req is None:
            continue

        for r in t.rows[1:]:
            cells = r.cells
            if idx_case >= len(cells) or idx_req >= len(cells):
                continue
            cid = re.sub(r"\s+", "", cells[idx_case].text or "")
            if not cid:
                continue
            mapped = case_to_reqs.get(cid)
            if not mapped:
                continue
            existing = set(REQ_ID_RE.findall(cells[idx_req].text or ""))
            merged = sorted(existing | set(mapped))
            if merged:
                _set_cell_text(cells[idx_req], ", ".join(merged))

    # 2) 追加“附录：需求追溯矩阵（自动补全）”
    doc.add_page_break()
    doc.add_heading("附录：需求追溯矩阵（自动补全）", level=1)
    doc.add_paragraph(f"生成日期：{date.today()}。说明：基于《1.1单元测试用例_修改版》与 SRS 自动汇总；未覆盖项需补充用例或在评审中裁剪说明。")

    # 反向映射：rid -> cases
    rid_to_cases: dict[str, list[str]] = {rid: [] for rid in sorted(reqs.keys())}
    for cid, rset in case_to_reqs.items():
        for rid in rset:
            if rid in rid_to_cases:
                rid_to_cases[rid].append(cid)

    # 输出表
    t = doc.add_table(rows=1 + len(rid_to_cases), cols=4)
    t.style = "Table Grid"
    t.cell(0, 0).text = "需求编号"
    t.cell(0, 1).text = "需求描述"
    t.cell(0, 2).text = "单元测试用例"
    t.cell(0, 3).text = "状态"
    for i, rid in enumerate(sorted(rid_to_cases.keys()), start=1):
        cases = sorted(set(rid_to_cases[rid]))
        t.cell(i, 0).text = rid
        t.cell(i, 1).text = reqs[rid].desc
        t.cell(i, 2).text = ", ".join(cases) if cases else "（待补充）"
        t.cell(i, 3).text = "已覆盖" if cases else "未覆盖"

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def main() -> None:
    if not SRC_11.exists():
        raise FileNotFoundError(SRC_11)
    if not SRC_11_MOD.exists():
        raise FileNotFoundError(SRC_11_MOD)

    srs = _find_srs_doc()
    reqs = extract_requirements_from_srs(srs)
    case_to_reqs = extract_trace_from_modified_cases(SRC_11_MOD)

    out = ROOT / f"1.1单元测试用例_追溯补全_{date.today().strftime('%Y%m%d')}.docx"
    update_original_doc(SRC_11, case_to_reqs, reqs, out)
    print(f"SRS: {srs}")
    print(f"Requirements: {len(reqs)}")
    print(f"Cases with trace: {len(case_to_reqs)}")
    print(f"Wrote: {out}")


if __name__ == "__main__":
    main()

