# -*- coding: utf-8 -*-
"""删除产品技术要求 docx 附录 A 中被遮盖的旧版服务端描述段落。"""
from __future__ import annotations

import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

DOCX_PATH = Path(r"C:\Users\lie76\Desktop\三维重建软件文档\产品技术要求0603.docx")

# 整段删除（a~d）
PARA_DELETE_KEYWORDS = (
    "将患者检查信息通过DICOM服务发送到图像设备",
    "将获取的DICOM格式图像文件进行无损保存",
    "支持DICOM WADO服务的接口",
    "提供数据存储功能，图像数据按序号存储在不同的位置",
)

# 第 205 段内仅删除含此句的 run（保留同段体系结构图）
MIXED_PARA_KEYWORD = "磁共振影像数据处理软件由系统服务和客户端组成"


def _run_text(run) -> str:
    return "".join(t.text or "" for t in run.findall(".//" + qn("w:t")))


def _remove_paragraph(paragraph) -> bool:
    element = paragraph._element
    parent = element.getparent()
    if parent is None:
        return False
    parent.remove(element)
    return True


def _remove_matching_runs(paragraph, keyword: str) -> int:
    removed = 0
    for run in list(paragraph._element.findall(qn("w:r"))):
        if keyword in _run_text(run):
            run.getparent().remove(run)
            removed += 1
    return removed


def main() -> None:
    if not DOCX_PATH.is_exists() if hasattr(DOCX_PATH, "is_exists") else not DOCX_PATH.exists():
        raise FileNotFoundError(DOCX_PATH)

    backup = DOCX_PATH.with_name(
        DOCX_PATH.stem + f"_backup_{datetime.now().strftime('%Y%m%d_%H%M%S')}" + DOCX_PATH.suffix
    )
    shutil.copy2(DOCX_PATH, backup)
    print(f"backup: {backup}")

    doc = Document(str(DOCX_PATH))
    deleted_paras = 0
    deleted_runs = 0

    # 先处理混合段落（含体系结构图），避免先删其它段落后索引变化
    for para in doc.paragraphs:
        if MIXED_PARA_KEYWORD in para.text:
            deleted_runs += _remove_matching_runs(para, MIXED_PARA_KEYWORD)
            print(f"removed run(s) in mixed diagram paragraph: {deleted_runs}")

    # 再整段删除 a~d（倒序遍历，避免 remove 后引用失效）
    for para in reversed(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        if any(k in text for k in PARA_DELETE_KEYWORDS):
            _remove_paragraph(para)
            deleted_paras += 1
            print(f"removed paragraph: {text[:60]}...")

    doc.save(str(DOCX_PATH))
    print(f"done: deleted {deleted_paras} paragraph(s), {deleted_runs} run(s)")
    print(f"saved: {DOCX_PATH}")


if __name__ == "__main__":
    main()
