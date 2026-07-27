#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
按《软件需求规格说明书_草稿_20260421.docx》模板格式生成《软件需求规格说明书》。

输出文件直接写入：
  C:\\Users\\lie76\\Desktop\\三维重建软件文档\\过审材料\\软件需求规格说明书_完善版_草稿.docx
"""

from __future__ import annotations

from dataclasses import dataclass
from datetime import datetime
from pathlib import Path

import docx


TEMPLATE = Path(r"C:\Users\lie76\Desktop\三维重建软件文档\过审材料\软件需求规格说明书_草稿_20260421.docx")
OUT = Path(r"C:\Users\lie76\Desktop\三维重建软件文档\过审材料\软件需求规格说明书_完善版_草稿.docx")


@dataclass(frozen=True)
class Req:
    rid: str
    module: str
    desc: str
    accept: str
    priority: str
    oref: str
    dref: str
    testcase: str = "—"
    evidence: str = "—"


@dataclass(frozen=True)
class DataReq:
    rid: str
    obj: str
    req: str
    accept: str
    priority: str
    oref: str
    dref: str
    testcase: str = "—"
    evidence: str = "—"


@dataclass(frozen=True)
class Nfr:
    rid: str
    topic: str
    req: str
    verify: str
    priority: str
    oref: str
    dref: str
    testcase: str = "—"
    evidence: str = "—"


def _set_paragraph_text(p, text: str) -> None:
    for r in p.runs:
        r.text = ""
    if p.runs:
        p.runs[0].text = text
    else:
        p.add_run(text)


def _reset_table_keep_header(tbl: docx.table.Table) -> None:
    while len(tbl.rows) > 1:
        tbl._tbl.remove(tbl.rows[1]._tr)  # noqa: SLF001 (python-docx internal)


def _append_row(tbl: docx.table.Table, cells: list[str]) -> None:
    row = tbl.add_row()
    for i, v in enumerate(cells):
        row.cells[i].text = v


def build() -> tuple[list[Req], list[DataReq], list[Nfr], list[Nfr], list[Nfr]]:
    fr: list[Req] = [
        Req(
            "FR-001",
            "用户登录",
            "用户能够通过用户名/密码登录系统，进入主工作界面。",
            "输入合法账号密码后进入主界面；错误密码提示失败且不登录；账户锁定时提示并拒绝登录。",
            "高",
            "概要：功能模块概要设计/登录模块",
            "详细：用户系统/认证流程",
        ),
        Req(
            "FR-002",
            "权限控制",
            "系统根据角色控制功能入口与操作权限（如用户管理仅管理员可用）。",
            "普通用户看不到或无法进入用户管理入口；越权操作被拒绝并提示。",
            "高",
            "概要：系统维护、运维与安全/权限模型",
            "详细：用户系统/安全模型",
        ),
        Req(
            "FR-003",
            "账户安全",
            "系统具备登录失败次数统计与自动锁定能力，防止暴力尝试。",
            "同一用户连续密码错误达到阈值后账户被锁定；锁定后无法登录；管理员可解锁/重置。",
            "高",
            "概要：系统出错处理设计/出错信息与补救措施",
            "详细：用户系统/认证流程（失败锁定）",
        ),
        Req(
            "FR-004",
            "修改密码",
            "用户在已登录状态下可修改密码；需校验原密码与新密码规则。",
            "原密码正确且新密码满足规则时修改成功；原密码错误提示失败且不修改。",
            "高",
            "概要：设置模块/用户设置",
            "详细：用户系统/密码存储",
        ),
        Req(
            "FR-005",
            "用户管理",
            "管理员可新增/编辑/禁用/删除用户，并可重置密码、锁定/解锁账户。",
            "管理员操作成功后用户列表与状态更新；重启后状态保持；非法操作提示原因。",
            "中",
            "概要：系统维护、运维与安全/用户数据维护",
            "详细：用户系统/界面类/与主窗口集成",
        ),
        Req(
            "FR-010",
            "语言切换",
            "系统支持中英文界面切换，默认中文；切换需确认且立即生效。",
            "在设置中切换语言后，主要界面文案切换并保持到下次启动；关键提示可正确汉化。",
            "中",
            "概要：设置模块/语言设置",
            "详细：主窗口/对外行为（中文提示）",
        ),
        Req(
            "FR-011",
            "主题与品牌化",
            "系统支持按配置启用/禁用主题切换；主题变化应同步到主窗口视觉。",
            "允许主题切换时可切换并生效；禁止时不提供入口或切换无效。",
            "低",
            "概要：程序系统结构/技术选型（QSS）",
            "详细：主题（ThemeSync + BrandingPreferences）",
        ),
        Req(
            "FR-020",
            "DICOM 导入",
            "系统支持通过 DICOM 目录导入影像并在多视图中展示。",
            "选择有效 DICOM 目录导入成功；影像在切片视图可浏览；导入过程有成功/失败反馈。",
            "高",
            "概要：影像导入与处理流程（概述）",
            "详细：影像导入与校验（主成功路径）",
        ),
        Req(
            "FR-021",
            "导入校验与回滚",
            "进入场景的新体数据必须通过统一规则校验；不通过则提示并回滚（移除节点）。",
            "导入不合规数据时给出原因；不合规节点不进入后续模块；场景无残留脏数据。",
            "高",
            "概要：数据一致性与恢复能力",
            "详细：影像导入校验/失败语义",
        ),
        Req(
            "FR-030",
            "影像查看",
            "系统提供多视图布局与基础查看交互（平移/缩放/窗宽窗位等按基线）。",
            "导入体数据后可在多视图中浏览；常用交互无明显卡顿；关键工具提示正确。",
            "高",
            "概要：用户接口（人机交互）/主工作区",
            "详细：主窗口/对外行为",
        ),
        Req(
            "FR-040",
            "医生批注",
            "系统提供医生批注能力，与当前选中体数据绑定，随场景保存。",
            "切换选中体数据时批注内容随之切换；保存并重开场景后批注仍存在。",
            "中",
            "概要：影像导入与处理流程（概述）",
            "详细：医生批注/绑定模型",
        ),
        Req(
            "FR-050",
            "成果保存与导出",
            "系统支持保存工作成果（场景保存为 MRB 或目录等），必要时支持模型导出。",
            "执行保存后生成可复打开的成果文件；导出失败给出错误提示并记录日志。",
            "高",
            "概要：外部接口/成果输出",
            "详细：医生标注与保存（主成功路径）",
        ),
        Req(
            "FR-060",
            "日志记录",
            "系统对关键操作与异常进行日志记录，便于追溯与排障。",
            "导入失败/认证失败等事件可在日志中定位到时间、级别与原因；日志可落盘并轮转（按配置）。",
            "高",
            "概要：异常处理/日志维护",
            "详细：日志（AppLogger）",
        ),
        Req(
            "FR-070",
            "授权检查（如启用）",
            "系统支持启动时授权校验；授权无效时提示并阻断或进入受限模式（按产品策略）。",
            "启用检查时：无效授权可识别并提示；有效授权可正常进入主流程。",
            "中",
            "概要：接口设计/内部接口（授权服务）",
            "详细：授权系统/扩展方式（思路）",
        ),
        Req(
            "FR-080",
            "壳层净化与模块白名单",
            "系统应隐藏/禁用与产品定位无关的原生入口，模块列表按白名单展示。",
            "主界面不暴露无关入口；模块下拉仅展示规定模块；不影响核心工作流。",
            "中",
            "概要：用户接口（人机交互）/壳层净化",
            "详细：主窗口/壳层净化策略",
        ),
        Req(
            "FR-090",
            "配置持久化",
            "系统应持久化基础配置（语言、主题、DICOM 路径等），并在下次启动生效。",
            "修改配置后重启仍保持；默认配置随安装包提供且可被用户配置覆盖。",
            "中",
            "概要：持久化设计/软件配置",
            "详细：SQLite 与 QSettings 的分工",
        ),
        Req(
            "FR-100",
            "退出与会话结束",
            "系统支持退出登录与退出应用；退出前按策略处理未保存内容并保证日志落盘。",
            "退出登录后回到登录界面且会话清理；退出应用后无残留进程；日志文件完整。",
            "中",
            "概要：系统退出",
            "详细：系统退出",
        ),
    ]

    dr: list[DataReq] = [
        DataReq(
            "DR-001",
            "影像数据",
            "支持 DICOM 影像数据输入与解析（元数据与像素数据由 DICOM/Slicer 机制承载）。",
            "可导入并展示；基本元数据可用于流程选择与校验。",
            "高",
            "概要：外部接口/影像输入",
            "详细：DICOM 导入后的对象映射",
        ),
        DataReq(
            "DR-002",
            "用户数据",
            "用户账号、角色、状态应持久化；口令不存明文，仅存不可逆摘要。",
            "重启后账号与状态保持；数据库中不出现明文口令；锁定策略可复现。",
            "高",
            "概要：持久化设计/用户与权限数据",
            "详细：用户系统/密码存储",
        ),
        DataReq(
            "DR-003",
            "配置数据",
            "默认配置与用户覆盖配置分层管理（安装默认 + 用户偏好），并可恢复到默认。",
            "删除/重置用户配置后回到默认；关键配置项重启后生效。",
            "中",
            "概要：配置维护/配置项清单",
            "详细：配置分层（设计原则）",
        ),
        DataReq(
            "DR-004",
            "日志数据",
            "日志落盘并按日期/策略滚动；不记录敏感像素，仅记录必要元信息与错误码。",
            "日志目录可定位；内容可追溯；抽查不包含像素数据与不必要敏感字段。",
            "中",
            "概要：日志维护/日志字段模板",
            "详细：系统维护与安全/安全与隐私",
        ),
        DataReq(
            "DR-005",
            "批注数据",
            "医生批注等轻量业务字段随 MRML 场景序列化保存，不写入关系库。",
            "保存并重开场景后批注存在；关系库不包含批注大对象。",
            "中",
            "概要：持久化设计/场景内业务数据",
            "详细：医生批注/绑定模型",
        ),
    ]

    pr: list[Nfr] = [
        Nfr(
            "PR-001",
            "启动时间",
            "送检基线硬件上冷启动到主界面可操作时间满足验收基线（建议≤10s，可按基线调整）。",
            "计时验证（基线机器）。",
            "中",
            "概要：性能目标（可量化模板）",
            "详细：启动阶段：阶段划分与时序",
        ),
        Nfr(
            "PR-002",
            "交互响应",
            "常用交互（切片滚动、平移缩放、窗宽窗位）响应在可接受范围内（建议<200ms/主观无卡顿）。",
            "人工体验 + 埋点（可选）。",
            "中",
            "概要：性能、容量与可靠性",
            "详细：性能与资源（落实说明）",
        ),
        Nfr(
            "PR-003",
            "导入耗时",
            "典型数据集导入耗时可接受（例如≤60s，按基线数据集定义）。",
            "用例测试（指定数据集、记录耗时）。",
            "中",
            "概要：测试与验证策略",
            "详细：影像导入与校验（主成功路径）",
        ),
    ]

    sr: list[Nfr] = [
        Nfr(
            "SR-001",
            "身份鉴别",
            "访问受控功能前必须完成登录鉴别；未登录不得进入受控工作流。",
            "未登录时受控入口不可用或跳转登录；用例验证。",
            "高",
            "概要：登录模块",
            "详细：用户系统/安全模型",
        ),
        Nfr(
            "SR-002",
            "口令安全",
            "口令以不可逆摘要存储；不得以明文或可逆方式持久化；不得在日志中输出明文口令。",
            "检查数据库/日志与配置：不含明文口令；用例/检查验证。",
            "高",
            "概要：设置模块/密码保存要求",
            "详细：用户系统/密码存储（SHA-256）",
        ),
        Nfr(
            "SR-003",
            "日志脱敏",
            "日志不记录敏感影像像素与不必要的患者敏感信息；必要标识采用脱敏/摘要。",
            "抽查日志内容；符合脱敏要求。",
            "高",
            "概要：安全与隐私设计",
            "详细：系统维护与安全/安全与隐私",
        ),
    ]

    rr: list[Nfr] = [
        Nfr(
            "RR-001",
            "异常处理",
            "可预期异常应提示并保持可继续操作；不可恢复错误应安全退出并保留日志证据。",
            "异常用例验证：提示清晰、系统可继续或安全退出；日志包含原因。",
            "中",
            "概要：系统出错处理设计",
            "详细：错误分类/与日志的互补",
        ),
        Nfr(
            "RR-002",
            "导入失败回滚",
            "导入或校验失败时不污染场景数据，保证主路径一致性。",
            "导入失败后场景无残留节点；再次导入可正常工作；用例验证。",
            "高",
            "概要：数据一致性与恢复能力",
            "详细：影像导入校验/失败语义",
        ),
        Nfr(
            "RR-003",
            "退出策略",
            "退出前按策略处理未保存内容（如适用），并确保日志完整落盘。",
            "退出用例验证：提示/处理符合策略；日志完整。",
            "中",
            "概要：系统退出",
            "详细：系统退出",
        ),
    ]

    return fr, dr, pr, sr, rr


def fill() -> None:
    doc = docx.Document(str(TEMPLATE))

    # Update narrative paragraphs (keep template sectioning/styles).
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t.startswith("本文档用于明确"):
            _set_paragraph_text(
                p,
                "本文档用于明确“医学影像三维重建软件（基于 3D Slicer 的桌面端应用）”的功能需求与非功能需求，"
                "作为概要设计、详细设计、实现、测试与验收的依据，并为送检材料提供可追溯的需求基线。",
            )
        elif t.startswith("本软件面向医学影像数据的导入"):
            _set_paragraph_text(
                p,
                "本软件面向医学影像数据的导入、查看、处理与三维重建，支持医生/操作员完成影像浏览、导入合规校验、"
                "三维结果生成与导出，并提供用户登录、权限控制、配置持久化与日志记录能力。"
                "影像与标注等业务数据以 MRML 场景为中心组织，账户与权限等结构化数据使用本地 SQLite 持久化。",
            )
        elif t.startswith("随着医学影像检查的普及"):
            _set_paragraph_text(
                p,
                "随着医学影像检查的普及，临床对影像三维可视化与结构分析需求增加。"
                "本软件用于在 Windows 桌面端对 DICOM 影像进行导入、校验、处理与三维重建，"
                "提升医生对结构的理解效率，并为术前评估/科研等提供辅助。"
                "软件采用 C++/Qt 外壳 + Slicer/VTK 能力 + Python 脚本化配置的组合形态。",
            )
        elif t.startswith("本软件功能包括但不限于"):
            _set_paragraph_text(
                p,
                "本软件功能包括但不限于：DICOM 导入与校验、影像多视图查看、三维重建与可视化、医生批注（如适用）、"
                "成果保存/导出、用户与权限、系统配置（语言/主题/路径等）与日志记录，以及对原生界面入口的壳层净化与白名单控制。",
            )

    # date paragraph: replace the first short 'yyyy年m月d日'
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t and "年" in t and "月" in t and "日" in t and len(t) <= 20:
            today = datetime.now()
            _set_paragraph_text(p, f"{today.year}年{today.month}月{today.day}日")
            break

    fr, dr, pr, sr, rr = build()

    t0, t1, t2, t3, t4, t5, t6, t7, t8, t9 = doc.tables[:10]

    _reset_table_keep_header(t0)
    for a, b, c in [
        ("DICOM", "医学数字成像与通信标准", "影像数据输入"),
        ("MRML", "Medical Reality Markup Language（Slicer 场景模型）", "场景/节点组织"),
        ("3D Slicer", "医学影像处理与三维可视化平台", "本软件基于其二次开发"),
        ("VTK", "Visualization Toolkit", "三维渲染与可视化"),
        ("Qt", "跨平台 GUI 框架（Widgets）", "主界面与事件循环"),
        ("QSettings", "Qt 配置存储机制", "用户偏好与默认配置覆盖"),
        ("SQLite", "轻量关系型数据库", "用户/权限等结构化数据持久化"),
    ]:
        _append_row(t0, [a, b, c])

    _reset_table_keep_header(t1)
    for r in [
        ("管理员", "账号与权限管理、配置维护", "创建/禁用用户、重置密码、解锁、配置维护"),
        ("医生/操作员", "影像处理与三维重建、批注与导出", "导入影像、查看、批注、保存/导出成果"),
    ]:
        _append_row(t1, list(r))

    _reset_table_keep_header(t2)
    for a, b in [
        ("操作系统", "Windows 10/11 64 位"),
        ("运行形态", "桌面端单机部署，离线可运行（无强制外部服务依赖）"),
        ("安装路径", "建议不包含中文与特殊字符（兼容性约束）；DICOM/缓存路径应可配置"),
        ("硬件", "满足常规影像处理与三维渲染需求（送检按版本基线给出配置）"),
    ]:
        _append_row(t2, [a, b])

    _reset_table_keep_header(t3)
    for a, b in [
        ("用户与权限", "登录、改密、角色权限、管理员用户管理"),
        ("影像导入", "DICOM 导入、完整性/规则校验、失败回滚"),
        ("影像查看", "多视图显示、常用交互与显示调整（按基线）"),
        ("三维重建", "重建流程入口、结果显示与导出（按版本范围）"),
        ("医生批注", "与当前选中体数据绑定，随场景保存"),
        ("成果保存/导出", "场景保存（MRB/目录等）、导出（如模型）"),
        ("配置与个性化", "语言、主题（可选）、DICOM 路径等配置持久化"),
        ("日志与诊断", "关键操作/异常日志落盘、轮转与保留策略"),
        ("壳层净化", "隐藏无关入口、模块白名单、减少误操作"),
        ("授权（可选）", "启动授权检查与受限策略（按产品策略）"),
    ]:
        _append_row(t3, [a, b])

    _reset_table_keep_header(t4)
    for r in fr:
        _append_row(t4, [r.rid, r.module, r.desc, r.accept, r.priority])

    _reset_table_keep_header(t5)
    for r in dr:
        _append_row(t5, [r.rid, r.obj, r.req, r.accept, r.priority])

    _reset_table_keep_header(t6)
    for r in pr:
        _append_row(t6, [r.rid, r.topic, r.req, r.verify, r.priority])

    _reset_table_keep_header(t7)
    for r in sr:
        _append_row(t7, [r.rid, r.topic, r.req, r.verify, r.priority])

    _reset_table_keep_header(t8)
    for r in rr:
        _append_row(t8, [r.rid, r.topic, r.req, r.verify, r.priority])

    _reset_table_keep_header(t9)
    for r in fr:
        _append_row(t9, [r.rid, r.oref, r.dref, r.testcase, r.evidence])
    for r in dr:
        _append_row(t9, [r.rid, r.oref, r.dref, r.testcase, r.evidence])
    for r in pr + sr + rr:
        _append_row(t9, [r.rid, r.oref, r.dref, r.testcase, r.evidence])

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))


if __name__ == "__main__":
    fill()

