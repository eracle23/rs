#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
参考：
  - C:\\Users\\lie76\\Desktop\\三维重建软件文档\\1.1单元测试用例.wps（已转为 docx 后分析其版式）
  - C:\\Users\\lie76\\Desktop\\三维重建软件文档\\1.2单元测试报告.doc（已转为 docx 后分析其版式）

生成两份互相关联的文档（.docx）：
  - 单元测试用例_Radiance_草稿_V1.docx
  - 单元测试报告_Radiance_草稿_V1.docx

两份文档通过 UT-xxx 用例编号关联，且在用例中标注需求追溯（FR/DR/PR/SR/RR）。
"""

from __future__ import annotations

from dataclasses import dataclass
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


OUT_DIR = Path(r"C:\Users\lie76\Desktop\三维重建软件文档\过审材料")
CASE_DOC = "单元测试用例_Radiance_草稿_V4.docx"
REPORT_DOC = "单元测试报告_Radiance_草稿_V4.docx"


@dataclass(frozen=True)
class UnitCase:
    cid: str
    name: str
    req: str
    module: str
    func_proto: str
    category: str  # 边界值分析 / 等价类划分 / 语句分支覆盖 / 异常分支
    func_desc: str
    inputs: str
    outputs: str
    stubs: str
    tester: str = ""
    result: str = "通过（模拟）"


def _title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(16)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _meta(doc: Document, rows: list[tuple[str, str]]) -> None:
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = "Table Grid"
    for i, (k, v) in enumerate(rows):
        t.cell(i, 0).text = k
        t.cell(i, 1).text = v


def _table(doc: Document, headers: list[str], data: list[list[str]]) -> None:
    t = doc.add_table(rows=1 + len(data), cols=len(headers))
    t.style = "Table Grid"
    for j, h in enumerate(headers):
        t.cell(0, j).text = h
    for i, row in enumerate(data, start=1):
        for j, cell in enumerate(row):
            t.cell(i, j).text = cell


def _proto_table(doc: Document, proto: str, desc: str, inp: str, out: str, ret: str) -> None:
    # 参考 1.1 模板：6x2，左列为“原型/描述/输入/输出/返回值/备注”
    t = doc.add_table(rows=6, cols=2)
    t.style = "Table Grid"
    pairs = [
        ("原型", proto),
        ("描述", desc),
        ("输入", inp),
        ("输出", out),
        ("返回值", ret),
        ("备注", "—"),
    ]
    for i, (k, v) in enumerate(pairs):
        t.cell(i, 0).text = k
        t.cell(i, 1).text = v


def _case_table(doc: Document, cases: list[UnitCase]) -> None:
    # 参考 1.1 模板：用例编号 | 输入 | 输出 | 桩函数描述
    _table(
        doc,
        ["用例编号", "输入（包括全程变量和输入参数）", "输出（包括全程变量和输出参数）", "桩函数描述"],
        [[c.cid, c.inputs, c.outputs, c.stubs] for c in cases],
    )


def build_cases() -> list[UnitCase]:
    # 以“可单测、可隔离”的逻辑为主：不要求启动完整 UI，不依赖真实 DICOM 数据库。
    return [
        # --- UserManager: hash/password/auth state ---
        UnitCase(
            "UT-001",
            "口令摘要：hashPassword 输出稳定",
            "SR-002/DR-002",
            "UserManager",
            "QString hashPassword(const QString& password)",
            "等价类划分",
            "对输入口令计算 SHA-256 十六进制摘要，用于持久化存储与比对。",
            "password='abc123'",
            "返回 64 字符十六进制字符串；同输入多次调用结果一致。",
            "无桩；纯函数单测。",
        ),
        UnitCase(
            "UT-002",
            "口令摘要：Unicode/中文口令编码一致性",
            "SR-002",
            "UserManager",
            "QString hashPassword(const QString& password)",
            "边界值分析",
            "对 UTF-8 编码口令计算摘要，避免中文口令在不同环境下不一致。",
            "password='医影123'（含中文）",
            "返回 64 字符十六进制；同输入一致；不抛异常。",
            "无桩；纯函数单测。",
        ),
        UnitCase(
            "UT-002A",
            "口令摘要：空字符串输入处理",
            "SR-002",
            "UserManager",
            "QString hashPassword(const QString& password)",
            "边界值分析",
            "空口令在实现层仍应有确定摘要输出（不用于放行登录，只验证算法稳定性）。",
            "password=''",
            "返回 64 字符十六进制；不抛异常。",
            "无桩；纯函数单测。",
        ),
        UnitCase(
            "UT-003",
            "认证失败计数：failed_attempts 自增",
            "FR-003/DR-002",
            "UserManager",
            "bool authenticate(const QString& username, const QString& password)",
            "语句分支覆盖",
            "密码错误时 failed_attempts 自增并写回数据库。",
            "用户存在；输入错误 password；初始 failed_attempts=0",
            "返回 false；failed_attempts=1；lastError 为“用户名或密码错误（剩余尝试次数: 4）”或等价文案。",
            "数据库桩：用内存 SQLite（QSQLITE :memory:）或 mock QSqlQuery。",
        ),
        UnitCase(
            "UT-003A",
            "认证失败：用户名不存在",
            "FR-001",
            "UserManager",
            "bool authenticate(const QString& username, const QString& password)",
            "等价类划分",
            "用户名不存在时应拒绝并提示“用户名或密码错误”。",
            "users 表中不存在 username；输入任意 password",
            "返回 false；lastError 为“用户名或密码错误”（或等价）。",
            "数据库桩。",
        ),
        UnitCase(
            "UT-003B",
            "认证阻断：账户禁用 is_active=0",
            "FR-002/DR-002",
            "UserManager",
            "bool authenticate(const QString& username, const QString& password)",
            "语句分支覆盖",
            "账户被禁用时阻断登录。",
            "用户存在；is_active=0；输入正确口令",
            "返回 false；lastError 提示“账户已被禁用”（或等价）。",
            "数据库桩。",
        ),
        UnitCase(
            "UT-003C",
            "认证阻断：账户已锁定 is_locked=1",
            "FR-003",
            "UserManager",
            "bool authenticate(const QString& username, const QString& password)",
            "语句分支覆盖",
            "账户锁定时无论口令正确与否都应拒绝登录。",
            "用户存在；is_locked=1；输入正确口令",
            "返回 false；lastError 提示“账户已被锁定”（或等价）。",
            "数据库桩。",
        ),
        UnitCase(
            "UT-004",
            "认证锁定：达到阈值后 is_locked=1",
            "FR-003",
            "UserManager",
            "bool authenticate(const QString& username, const QString& password)",
            "边界值分析",
            "连续错误达到阈值（当前实现为 5 次）锁定账户。",
            "同一用户连续 5 次错误口令",
            "第 5 次返回 false；is_locked=1；lastError 提示锁定；触发 userLocked 信号（如可观测）。",
            "数据库桩 + 信号观察桩（QSignalSpy）。",
        ),
        UnitCase(
            "UT-005",
            "认证成功：last_login 更新且 failed_attempts 清零",
            "FR-001/FR-003",
            "UserManager",
            "bool authenticate(const QString& username, const QString& password)",
            "语句分支覆盖",
            "认证成功后更新 last_login 并清零失败次数。",
            "用户存在；输入正确口令；failed_attempts>0",
            "返回 true；last_login 更新为 ISO 时间串；failed_attempts=0。",
            "数据库桩（内存 SQLite）+ 时间可接受范围断言。",
        ),
        UnitCase(
            "UT-006",
            "修改密码：原密码错误阻断",
            "FR-004",
            "UserManager",
            "bool changePassword(int userId, const QString& oldPassword, const QString& newPassword)",
            "语句分支覆盖",
            "原密码错误时不更新 password_hash。",
            "oldPassword 错误，newPassword 合法",
            "返回 false；lastError='原密码错误'；password_hash 不变。",
            "数据库桩。",
        ),
        UnitCase(
            "UT-006A",
            "修改密码：成功路径更新 password_hash",
            "FR-004",
            "UserManager",
            "bool changePassword(int userId, const QString& oldPassword, const QString& newPassword)",
            "等价类划分",
            "原密码正确时应更新 password_hash。",
            "oldPassword 正确；newPassword 合法",
            "返回 true；password_hash 改变；可用新口令通过 authenticate。",
            "数据库桩。",
        ),
        UnitCase(
            "UT-007",
            "重置密码：解除锁定并清零失败次数",
            "FR-005/FR-003",
            "UserManager",
            "bool resetPassword(int userId, const QString& newPassword)",
            "语句分支覆盖",
            "管理员重置密码时解除锁定与失败计数，恢复可登录状态。",
            "用户 is_locked=1；failed_attempts=5",
            "返回 true；is_locked=0；failed_attempts=0；password_hash 更新为新摘要。",
            "数据库桩 + 信号观察（userUnlocked）。",
        ),
        UnitCase(
            "UT-008",
            "用户库初始化：建表成功且可重复调用",
            "DR-002",
            "UserManager",
            "bool initialize()",
            "语句分支覆盖",
            "打开/创建 users.db 并执行建表语句；重复初始化不应失败。",
            "可写目录；首次/重复调用 initialize()",
            "返回 true；users 表存在；重复调用仍返回 true。",
            "文件系统桩：临时目录映射 AppLocalDataLocation；SQLite 可用。",
        ),
        UnitCase(
            "UT-008A",
            "用户库初始化：数据库不可写路径失败语义",
            "RR-001/DR-002",
            "UserManager",
            "bool initialize()",
            "异常分支",
            "数据库路径不可写时应失败并提供可诊断 lastError。",
            "将 AppLocalDataLocation 指向只读目录",
            "返回 false；lastError 包含“无法打开用户数据库”或等价信息。",
            "文件系统桩：只读目录。",
        ),
        UnitCase(
            "UT-009",
            "默认管理员：首次创建、二次不重复插入",
            "FR-001/DR-002",
            "UserManager",
            "bool createDefaultAdmin()",
            "语句分支覆盖",
            "若不存在 admin 则插入默认管理员；已存在则直接返回成功。",
            "空 users 表→调用两次 createDefaultAdmin()",
            "第一次插入 admin；第二次不新增重复记录；均返回 true。",
            "数据库桩：内存 SQLite 或临时 users.db。",
        ),
        UnitCase(
            "UT-012",
            "新增用户：username 唯一约束冲突返回失败",
            "FR-005/DR-002",
            "UserManager",
            "bool addUser(const UserInfo& user)",
            "边界值分析",
            "username TEXT UNIQUE；重复用户名应失败并提供 lastError。",
            "已存在 username='u1'；再次 addUser(u1)",
            "返回 false；lastError 包含“添加用户失败”或唯一约束失败信息。",
            "数据库桩。",
        ),
        UnitCase(
            "UT-012A",
            "新增用户：空用户名/空口令摘要拒绝（如实现含校验）",
            "FR-001/FR-005",
            "UserManager",
            "bool addUser(const UserInfo& user)",
            "异常分支",
            "输入校验应阻断空用户名或空摘要（若当前未实现，可作为差距记录）。",
            "username='' 或 passwordHash=''",
            "返回 false 或触发断言；lastError 给出原因（或记录差距）。",
            "无桩/数据库桩。",
        ),
        UnitCase(
            "UT-013",
            "更新用户：变更 role/is_active 生效",
            "FR-002/FR-005",
            "UserManager",
            "bool updateUser(const UserInfo& user)",
            "等价类划分",
            "管理员调整用户角色与启用状态，更新应落库。",
            "user.id 有效；role='admin'；is_active=0",
            "返回 true；重新查询该用户字段已更新。",
            "数据库桩。",
        ),
        UnitCase(
            "UT-014",
            "删除用户：记录从 users 表移除",
            "FR-005",
            "UserManager",
            "bool deleteUser(int userId)",
            "等价类划分",
            "删除用户记录；失败应返回 false 并设置 lastError。",
            "userId 有效",
            "返回 true；查询不到该 id。",
            "数据库桩。",
        ),
        UnitCase(
            "UT-015",
            "锁定/解锁：is_locked 与 failed_attempts 按规则更新",
            "FR-003/FR-005",
            "UserManager",
            "bool lockUser(int userId) / bool unlockUser(int userId)",
            "语句分支覆盖",
            "锁定置 is_locked=1；解锁置 is_locked=0 且 failed_attempts=0。",
            "userId 有效；failed_attempts=3",
            "lockUser 返回 true 且 is_locked=1；unlockUser 返回 true 且 is_locked=0, failed_attempts=0。",
            "数据库桩 + 信号观察（userLocked/userUnlocked）。",
        ),
        UnitCase(
            "UT-016",
            "获取用户列表：按 id 排序返回稳定",
            "FR-005",
            "UserManager",
            "QVector<UserInfo> getAllUsers()",
            "语句分支覆盖",
            "查询用户列表并按 id 排序，便于管理界面稳定显示。",
            "插入多条用户 id 乱序",
            "返回列表按 id 升序。",
            "数据库桩。",
        ),
        UnitCase(
            "UT-016A",
            "按用户名查询：getUserByUsername 不存在返回空对象",
            "FR-005",
            "UserManager",
            "UserInfo getUserByUsername(const QString& username)",
            "等价类划分",
            "查询不到用户时应返回空 UserInfo（id=0 或默认值），并可设置 lastError（按实现）。",
            "username='not_exists'",
            "返回默认 UserInfo；不抛异常。",
            "数据库桩。",
        ),
        UnitCase(
            "UT-016B",
            "按 ID 查询：getUserById 不存在返回空对象",
            "FR-005",
            "UserManager",
            "UserInfo getUserById(int userId)",
            "等价类划分",
            "查询不到用户时返回空对象，不应崩溃。",
            "userId=999999（不存在）",
            "返回默认 UserInfo；不抛异常。",
            "数据库桩。",
        ),
        UnitCase(
            "UT-016C",
            "会话态：登录成功后 isUserLoggedIn=true",
            "FR-001",
            "UserManager",
            "bool isUserLoggedIn() / UserInfo getCurrentUser()",
            "语句分支覆盖",
            "登录成功后会话态应设置为已登录，并可获取当前用户信息。",
            "authenticate 成功一次",
            "isUserLoggedIn()=true；getCurrentUser().username 正确。",
            "数据库桩。",
        ),
        UnitCase(
            "UT-016D",
            "会话态：logout 清理 currentUser 并置 isUserLoggedIn=false",
            "FR-100",
            "UserManager",
            "void logout()",
            "语句分支覆盖",
            "退出登录应清理当前用户并恢复未登录态。",
            "先登录成功→调用 logout()",
            "isUserLoggedIn()=false；getCurrentUser() 为默认对象。",
            "无桩/数据库桩。",
        ),
        UnitCase(
            "UT-016E",
            "错误信息：getLastError 返回最近一次失败原因",
            "RR-001",
            "UserManager",
            "QString getLastError()",
            "语句分支覆盖",
            "失败后 lastError 可供 UI 展示与日志记录。",
            "触发一次认证失败/查询失败",
            "getLastError() 非空且与失败原因一致。",
            "数据库桩。",
        ),
        # --- Config / Branding / Theme ---
        UnitCase(
            "UT-010",
            "主题开关解析：AllowThemeSwitch 默认值",
            "FR-011/DR-003",
            "Theme/Branding",
            "bool themeSwitchAllowed()",
            "等价类划分",
            "从默认配置/用户配置读取 Radiance/AllowThemeSwitch，缺省按产品策略（通常 false）。",
            "配置缺省不含 AllowThemeSwitch",
            "返回 false（或产品默认）；不抛异常。",
            "配置桩：临时 QSettings scope（IniFormat + 临时文件）。",
        ),
        UnitCase(
            "UT-011",
            "主题开关解析：显式 true 生效",
            "FR-011",
            "Theme/Branding",
            "bool themeSwitchAllowed()",
            "等价类划分",
            "显式配置为 true 时允许主题切换。",
            "AllowThemeSwitch=true",
            "返回 true。",
            "配置桩。",
        ),
        UnitCase(
            "UT-017",
            "主题同步：环境变量禁用时不执行主题同步",
            "FR-011",
            "Theme/Branding",
            "ThemeSync::applyBranding(...)",
            "异常分支",
            "当环境变量禁用开关存在时，跳过主题同步以便现场排障。",
            "设置 YOURAPP_DISABLE_THEMESYNC=1",
            "applyBranding 不对窗口样式做修改（或直接返回）。",
            "环境变量桩 + 顶层 QWidget 样式断言（最小窗口）。",
        ),
        UnitCase(
            "UT-018",
            "语言设置：语言键持久化写入并可读取",
            "FR-010/DR-003",
            "Config",
            "QSettings read/write language key",
            "语句分支覆盖",
            "语言切换应写入用户配置并在下次启动读取。",
            "写入 Language='en-US' 后重新打开 QSettings",
            "读取到 en-US；无异常。",
            "QSettings 桩：IniFormat + 临时文件。",
        ),
        UnitCase(
            "UT-018A",
            "配置重置：删除用户配置后恢复默认值",
            "DR-003",
            "Config",
            "resetUserSettings()",
            "异常分支",
            "删除/重置用户配置文件后应回到默认配置（若实现提供重置入口）。",
            "先写入自定义配置→执行 reset→重启读取",
            "读取到默认值；无崩溃。",
            "文件系统桩 + QSettings 临时文件。",
        ),
        UnitCase(
            "UT-018B",
            "配置优先级：用户配置覆盖默认配置",
            "DR-003",
            "Config",
            "readSetting(key) with defaultSettings + userSettings",
            "语句分支覆盖",
            "同一键存在于默认配置与用户配置时，应以用户配置为准。",
            "defaultSettings: K=0；userSettings: K=1",
            "读取结果为 1。",
            "QSettings 双层桩（两个 ini 文件）。",
        ),
        # --- Import validation (rule decisions) ---
        UnitCase(
            "UT-020",
            "导入校验：不合规输入返回失败语义",
            "FR-021/RR-002",
            "VolumeImportValidator",
            "Validate(volume) -> (pass, reason)",
            "等价类划分",
            "对新加入场景的体数据执行规则校验，失败应给出原因并可触发回滚。",
            "构造触发失败的假 volume 元数据（来源/尺寸/灰度任一）",
            "pass=false；reason 非空且包含类别（来源/几何/灰度）提示。",
            "MRML/VTK 桩：构造最小 vtkMRMLScalarVolumeNode 或 mock 元数据接口。",
        ),
        UnitCase(
            "UT-021",
            "导入校验：合规输入返回通过",
            "FR-020/FR-021",
            "VolumeImportValidator",
            "Validate(volume) -> (pass, reason)",
            "等价类划分",
            "合规体数据应通过校验进入后续流程。",
            "构造合规 volume 元数据",
            "pass=true；reason 为空或为可忽略提示。",
            "同上。",
        ),
        UnitCase(
            "UT-022",
            "导入校验：来源层失败（非 DICOM/无存储信息）",
            "FR-021",
            "VolumeImportValidator",
            "Validate(volume) -> (pass, reason)",
            "语句分支覆盖",
            "来源层失败时应返回失败并提示来源不合规。",
            "volume 无有效 storage / 来源标记缺失",
            "pass=false；reason 包含“来源/非 DICOM”类提示。",
            "MRML 桩：构造无 storage 的 volume 或 mock 属性。",
        ),
        UnitCase(
            "UT-023",
            "导入校验：几何层失败（异常尺寸/spacing）",
            "FR-021/RR-002",
            "VolumeImportValidator",
            "Validate(volume) -> (pass, reason)",
            "边界值分析",
            "几何参数异常应被拒绝，避免内存/渲染风险。",
            "volume spacing/extent 超阈值",
            "pass=false；reason 包含“几何/尺寸/spacing”类提示。",
            "VTK/MRML 桩或 mock 几何读取接口。",
        ),
        UnitCase(
            "UT-024",
            "导入校验：灰度层失败（标量范围异常）",
            "FR-021",
            "VolumeImportValidator",
            "Validate(volume) -> (pass, reason)",
            "边界值分析",
            "标量范围不符合产品假设时应拒绝并提示。",
            "volume scalarRange 超阈值",
            "pass=false；reason 包含“灰度/范围”类提示。",
            "VTK 数据桩或 mock scalarRange。",
        ),
        UnitCase(
            "UT-024A",
            "导入校验：空指针/无效 volume 输入处理",
            "RR-001",
            "VolumeImportValidator",
            "Validate(volume) -> (pass, reason)",
            "异常分支",
            "volume 为空或必要字段缺失时应失败并给出 reason，不应崩溃。",
            "volume=None 或缺少标量范围/几何信息",
            "pass=false；reason 非空；无崩溃。",
            "mock 接口或空对象桩。",
        ),
        UnitCase(
            "UT-024B",
            "导入校验：多规则同时失败时原因优先级稳定",
            "FR-021",
            "VolumeImportValidator",
            "Validate(volume) -> (pass, reason)",
            "语句分支覆盖",
            "当来源+几何+灰度同时不合规时，reason 的优先级应稳定（便于用户理解与测试复现）。",
            "构造同时触发三类失败的 volume",
            "reason 按既定优先级输出（如先来源，再几何，再灰度）。",
            "VTK/MRML 桩。",
        ),
        # --- Logging policy ---
        UnitCase(
            "UT-030",
            "日志文件名：按日期轮转命名",
            "FR-060/DR-004",
            "AppLogger",
            "logFilePath(date) -> path",
            "等价类划分",
            "日志按日期轮转，文件名包含日期便于追溯。",
            "date=2026-04-30",
            "输出路径包含 '2026-04-30'（或等价格式），目录可创建。",
            "文件系统桩：临时目录；不写真实业务目录。",
        ),
        UnitCase(
            "UT-031",
            "日志清理：超过保留天数的旧日志被删除",
            "FR-060",
            "AppLogger",
            "cleanupOldLogs(retentionDays)",
            "边界值分析",
            "按策略清理旧日志，避免磁盘占用无限增长。",
            "retentionDays=7；准备 10 天旧日志文件",
            "超过 7 天的被删除；近 7 天保留；不误删当前文件。",
            "文件系统桩：临时目录 + 人工构造旧文件时间戳。",
        ),
        UnitCase(
            "UT-032",
            "日志目录回退：首选目录不可写时回退可写目录",
            "FR-060/RR-001",
            "AppLogger",
            "initLogDir(preferredDir) -> actualDir",
            "异常分支",
            "当首选日志目录不可写时，应回退到用户可写目录保证日志落盘。",
            "preferredDir=只读目录",
            "返回可写目录；写入日志不失败。",
            "文件系统桩：只读目录 + 临时目录。",
        ),
        UnitCase(
            "UT-033",
            "日志级别：低级别日志可被过滤（若实现支持）",
            "FR-060",
            "AppLogger",
            "setLogLevel(level) + log(level,msg)",
            "语句分支覆盖",
            "当设置过滤级别时，低于阈值的日志不应落盘（如实现支持）。",
            "level=WARNING；写 INFO 与 WARNING",
            "INFO 不落盘；WARNING 落盘（或按实现说明）。",
            "文件系统桩 + 读取文件内容断言。",
        ),
        UnitCase(
            "UT-033A",
            "日志脱敏：输出中不包含口令明文关键字（抽查）",
            "SR-003/SR-002",
            "AppLogger",
            "log(msg) content policy",
            "异常分支",
            "确保日志不记录口令明文（抽查规则）。",
            "写入包含 password=*** 的消息（模拟）",
            "日志中不应出现敏感明文字段（或必须脱敏）。",
            "文件系统桩 + 正则抽查。",
        ),
        UnitCase(
            "UT-034",
            "日志写入失败：磁盘满/不可写时不崩溃（最小容错）",
            "RR-001",
            "AppLogger",
            "writeLogLine(...)",
            "异常分支",
            "写日志失败时应避免崩溃；可降级为 qDebug 或静默失败（按实现）。",
            "模拟写入失败（只读目录/抛异常）",
            "不抛致命异常；主流程可继续。",
            "文件系统桩。",
        ),
        # --- Shell cleaner / i18n fallback ---
        UnitCase(
            "UT-040",
            "Tooltip 汉化匹配：去空白/句点归一化",
            "FR-010",
            "RadianceShellCleaner",
            "normalizeTooltip(text) -> normalized",
            "边界值分析",
            "对 tooltip 文案做 trim/去尾句点等归一化，确保运行期替换鲁棒。",
            "text='Show how ... plane. '（含尾句点与空格）",
            "normalized 不含尾句点与多余空白；可用于 key 匹配。",
            "无桩；纯函数或可抽取逻辑单测。",
        ),
        UnitCase(
            "UT-041",
            "Tooltip 汉化替换：命中时替换为中文",
            "FR-010",
            "RadianceShellCleaner",
            "localizeTooltip(widgetTree)",
            "等价类划分",
            "扫描窗口组件树并替换目标 tooltip。",
            "构造包含目标 tooltip 的假 QWidget 树",
            "tooltip 被替换为中文；非目标不变。",
            "Qt Widget 桩：最小 QWidget + setToolTip。",
        ),
        UnitCase(
            "UT-042",
            "Tooltip 汉化：未命中时不误替换其它 tooltip",
            "FR-010",
            "RadianceShellCleaner",
            "localizeTooltip(widgetTree)",
            "语句分支覆盖",
            "只替换目标 tooltip，避免误伤其它控件提示。",
            "构造多个控件 tooltip 相似但不相同",
            "仅目标被替换；其余保持原文。",
            "Qt Widget 桩。",
        ),
        UnitCase(
            "UT-043",
            "壳层净化：模块白名单过滤逻辑（若有独立函数）",
            "FR-080",
            "RadianceShellCleaner",
            "filterModules(allModules) -> allowed",
            "等价类划分",
            "模块列表按白名单过滤，降低误操作风险。",
            "allModules 含白名单与非白名单项",
            "返回仅包含白名单项；顺序稳定。",
            "无桩：纯容器逻辑或 mock 模块名列表。",
        ),
        UnitCase(
            "UT-044",
            "壳层净化：重复应用不会产生重复副作用（幂等）",
            "FR-080/RR-001",
            "RadianceShellCleaner",
            "apply(window) / applyRules(...)",
            "语句分支覆盖",
            "多次调用净化逻辑不应导致重复隐藏、重复连接信号或异常。",
            "对同一 window 连续调用 apply 两次",
            "UI 状态稳定；不出现重复条目/重复弹窗；无异常日志。",
            "Qt Widget 桩（最小 QMainWindow）。",
        ),
        UnitCase(
            "UT-045",
            "授权开关：EnableStartupCheck 解析缺省与显式值",
            "FR-070/DR-003",
            "License",
            "read License/EnableStartupCheck",
            "等价类划分",
            "授权检查开关缺省/显式值应解析一致，便于正式/研发模式切换。",
            "缺省不含该键；显式 true/false 各一次",
            "缺省按产品策略；显式值优先。",
            "QSettings 桩。",
        ),
    ]


def generate_case_doc(cases: list[UnitCase]) -> Path:
    doc = Document()
    _title(doc, "单元测试用例（Radiance）")
    _meta(
        doc,
        [
            ("文档版本", "V4.0 草稿"),
            ("编制日期", str(date.today())),
            ("软件名称", "医学影像三维重建软件（Radiance / 基于 3D Slicer）"),
            ("说明", "参考 1.1 单元测试用例模板：按“函数原型 → 用例表（输入/输出/桩函数）”编写；用例编号 UT-xxx 与《单元测试报告》一致。"),
        ],
    )

    doc.add_paragraph("")
    doc.add_heading("1. 概述", level=1)
    doc.add_paragraph(
        "单元测试关注可隔离、可重复的最小逻辑单元，采用桩/Mock 隔离数据库、文件系统、MRML/VTK 等外部依赖。"
        "本用例集覆盖用户与权限、配置解析、导入校验判定、日志策略、壳层汉化替换等关键逻辑。"
    )

    doc.add_heading("2. 用例清单", level=1)
    _table(
        doc,
        ["用例编号", "用例名称", "模块", "类型", "需求追溯"],
        [[c.cid, c.name, c.module, c.category, c.req] for c in cases],
    )

    doc.add_heading("3. 用例明细（按函数原型组织）", level=1)
    # group by func_proto
    groups: dict[str, list[UnitCase]] = {}
    for c in cases:
        groups.setdefault(c.func_proto, []).append(c)

    for proto, g in groups.items():
        doc.add_paragraph("")
        doc.add_heading(f"3.x {proto}", level=2)
        # prototype summary uses first case
        first = g[0]
        _proto_table(
            doc,
            proto=proto,
            desc=first.func_desc,
            inp="（见下表用例输入）",
            out="（见下表用例输出）",
            ret="按实现：bool/QString/结构体（以当前版本为准）",
        )
        doc.add_paragraph("")
        # 参考模板“边界值/等价类/语句覆盖”的组织方式
        for cat in ["边界值分析", "等价类划分", "语句分支覆盖", "异常分支"]:
            gg = [x for x in g if x.category == cat]
            if not gg:
                continue
            doc.add_heading(f"{cat}", level=3)
            _case_table(doc, gg)

    out = OUT_DIR / CASE_DOC
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return out


def generate_report_doc(cases: list[UnitCase]) -> Path:
    doc = Document()
    _title(doc, "单元测试报告（Radiance）")
    _meta(
        doc,
        [
            ("文档版本", "V4.0 草稿"),
            ("报告日期", str(date.today())),
            ("关联用例文档", "《单元测试用例_Radiance_草稿_V4》"),
            ("说明", "参考 1.2 单元测试报告模板：按模块汇总“用例编号/用例名称/测试人员/测试结果”。"),
        ],
    )

    doc.add_paragraph("")
    doc.add_heading("1. 测试概述", level=1)
    doc.add_paragraph(
        "本报告汇总单元测试执行情况。正式送检版应填写真实测试人员、日期、缺陷编号与证据（日志/截图/输出）。本草稿以“通过（模拟）”占位。"
    )

    doc.add_heading("2. 结果汇总", level=1)
    n = len(cases)
    passed = sum(1 for c in cases if c.result.startswith("通过"))
    _table(
        doc,
        ["统计项", "数量"],
        [
            ["计划用例数", str(n)],
            ["通过", str(passed)],
            ["不通过", "0（待填写）"],
            ["阻塞", "0（待填写）"],
            ["N/A", "0（待填写）"],
        ],
    )
    doc.add_paragraph("说明：本草稿以“通过（模拟）”占位；正式报告应填写真实执行人/日期/缺陷编号与证据。")

    doc.add_heading("2.1 测试人员信息", level=2)
    # 对齐模板常见“职务/姓名/岗位/单位/测试日期”
    _table(
        doc,
        ["职务", "姓名", "岗位", "单位", "测试日期"],
        [
            ["组长", "（填写）", "软件工程师", "研发部", str(date.today())],
            ["成员", "（填写）", "软件工程师", "研发部", str(date.today())],
            ["成员", "（填写）", "软件工程师", "研发部", str(date.today())],
        ],
    )

    # group by module
    mod: dict[str, list[UnitCase]] = {}
    for c in cases:
        mod.setdefault(c.module, []).append(c)

    doc.add_heading("3. 模块级用例执行明细", level=1)
    for m, g in mod.items():
        doc.add_paragraph("")
        # 模块表：首行“模块名称 | xxx”，第二行“测试时间 | ”，第三行 header
        t = doc.add_table(rows=3 + len(g), cols=4)
        t.style = "Table Grid"
        t.cell(0, 0).text = "模块名称"
        t.cell(0, 1).text = m
        t.cell(0, 2).text = ""
        t.cell(0, 3).text = ""
        t.cell(1, 0).text = "测试时间"
        t.cell(1, 1).text = str(date.today())
        t.cell(1, 2).text = ""
        t.cell(1, 3).text = ""
        t.cell(2, 0).text = "用例编号"
        t.cell(2, 1).text = "用例名称"
        t.cell(2, 2).text = "测试人员"
        t.cell(2, 3).text = "测试结果"
        for i, c in enumerate(g, start=3):
            t.cell(i, 0).text = c.cid
            t.cell(i, 1).text = c.name
            t.cell(i, 2).text = c.tester or "（填写）"
            t.cell(i, 3).text = c.result

    doc.add_heading("4. 缺陷与备注（如有）", level=1)
    doc.add_paragraph("本草稿：无缺陷。若出现不通过/阻塞，请在此登记缺陷编号与回归状态。")
    _table(
        doc,
        ["缺陷编号", "等级", "摘要", "关联用例", "状态", "备注"],
        [["—", "—", "无", "—", "—", "—"]],
    )

    doc.add_heading("5. 结论", level=1)
    doc.add_paragraph(
        "结论（草稿）：单元测试覆盖关键可隔离逻辑点，为集成测试与系统测试提供基础保障。正式报告需补齐真实执行记录与证据。"
    )

    out = OUT_DIR / REPORT_DOC
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return out


def main() -> None:
    cases = build_cases()
    out1 = generate_case_doc(cases)
    out2 = generate_report_doc(cases)
    print("Wrote:", out1)
    print("Wrote:", out2)


if __name__ == "__main__":
    main()

