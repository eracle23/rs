# -*- coding: utf-8 -*-
"""修订 CH3.5.5.11 系列网络安全文档，与医学影像三维重建软件（Vision Magic Ecosystem）对齐。"""
from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document

BASE = Path(r"C:\Users\lie76\Desktop\三维重建软件文档\过审材料\软件资料需求2026.06.23")

SOFTWARE_NAME = "医学影像三维重建软件"
SOFTWARE_EN = "Vision Magic Ecosystem"
EXE_NAME = "VisionMagicEcosystem.exe"
VERSION = "V1.0"
MODEL = "VisionMagicEcosystem"

# 22 项网络安全能力（研究报告 TABLE 3）
CAPABILITY_ROWS = [
    [
        "C01",
        "自动注销（ALOF）",
        "否",
        "本软件无自动超时注销；用户通过【退出登录】主动结束会话，已在说明书中说明。",
        "/",
    ],
    [
        "C02",
        "审核控制（AUDT）",
        "是",
        "/",
        "AppLogger 记录应用程序启动/关闭、用户操作及错误信息，日志存于 "
        "%APPDATA%\\VisionMagicEcosystem\\logs。",
    ],
    [
        "C03",
        "授权（AUTH）",
        "是",
        "/",
        "1.启动前须通过用户名/密码登录；2.单实例运行限制；"
        "3.可按部署策略启用授权/许可校验（如加密狗）；4.修改密码需验证原密码。",
    ],
    [
        "C04",
        "节点鉴别（NAUT）",
        "否",
        "本软件以单机本地运行为主，DICOM 以本地文件夹导入为主，不强制 DICOM 网络节点鉴别。",
        "/",
    ],
    [
        "C05",
        "人员鉴别（PAUT）",
        "是",
        "/",
        "1.区分管理员与普通用户；2.用户名 1～50 字符、密码不少于 6 位；"
        "3.连续 5 次登录失败锁定 15 分钟或管理员解锁；4.管理员可添加/锁定/解锁/重置密码。",
    ],
    [
        "C06",
        "连通性（CONN）",
        "否",
        "单机软件，正常运行不依赖外部网络连接。",
        "/",
    ],
    [
        "C07",
        "物理防护（PLOK）",
        "是",
        "/",
        "建议专用工作站部署；说明书中对物理访问与设备放置提出要求。",
    ],
    [
        "C08",
        "系统加固（SAHD）",
        "是",
        "/",
        "软件安装与升级由专业技术人员或授权人员完成。",
    ],
    [
        "C09",
        "数据去标识化与匿名化（DIDT）",
        "是",
        "/",
        "登录密码密文显示；临床使用数据脱敏要求见产品说明书。",
    ],
    [
        "C10",
        "数据完整性与真实性（IGAU）",
        "是",
        "/",
        "1.本地 DICOM 导入校验（VolumeImportValidator）；2.删除操作需确认；"
        "3.导入后可在视图中核对 DICOM 标签信息。",
    ],
    [
        "C11",
        "数据备份与灾难恢复（DTBK）",
        "是",
        "/",
        "支持场景/项目保存（.mrml/.mrb 等）及重新打开，异常退出后可通过保存文件恢复。",
    ],
    [
        "C12",
        "数据存储保密性与完整性（STCF）",
        "是",
        "/",
        "用户数据库与日志存于本机；建议操作系统账户权限控制，说明书有相关要求。",
    ],
    [
        "C13",
        "数据传输保密性（TXCF）",
        "否",
        "单机本地导入为主，不涉及强制性网络数据传输。",
        "/",
    ],
    [
        "C14",
        "数据传输完整性（TXIG）",
        "是",
        "/",
        "本地 DICOM 导入后可完整浏览切片，层数与源文件一致。",
    ],
    [
        "C15",
        "网络安全补丁升级（CSUP）",
        "否",
        "单机部署，操作系统补丁由用户/运维按说明书维护。",
        "/",
    ],
    [
        "C16",
        "现成软件清单（SBOM）",
        "是",
        "/",
        "产品含 3D Slicer/Qt/Python 等现成软件，提供现成软件清单。",
    ],
    [
        "C17",
        "现成软件维护（RDMP）",
        "是",
        "/",
        "对现成软件组件建立版本跟踪与维护策略。",
    ],
    [
        "C18",
        "网络安全使用指导（SGUD）",
        "是",
        "/",
        "产品说明书中规定网络安全注意事项与警告事项。",
    ],
    [
        "C19",
        "网络安全特性配置(CNFS)",
        "是",
        "/",
        "用户可修改自身密码；管理员可重置密码及管理账户。",
    ],
    [
        "C20",
        "紧急访问（EMRG）",
        "否",
        "本产品不包含紧急访问功能。",
        "/",
    ],
    [
        "C21",
        "远程访问与控制（RMOT）",
        "否",
        "本产品不适用远程访问与远程维护。",
        "/",
    ],
    [
        "C22",
        "恶意软件探测与防护（MLDP）",
        "是",
        "/",
        "说明书规定推荐运行环境与安全软件；已在火绒等杀毒软件共存环境下验证。",
    ],
]

COMM_DATA_ROWS = [
    [
        "DICOM 影像数据",
        "DICOM 本地导入/加载",
        "单向",
        "从本地磁盘读取 DICOM 文件并加载至软件进行浏览与三维重建",
    ],
    [
        "场景/项目数据",
        "保存与加载",
        "双向",
        "将当前工作场景保存至本地文件并在后续会话中重新打开",
    ],
]

INTERFACE_ROWS = [
    [
        "本地文件/DICOM",
        "使用者",
        "导入本地 DICOM 或影像文件",
        "加载医学影像并进行三维重建与浏览",
        "DICOM 影像/体数据",
        "格式：dcm、nrrd、nii 等；本地文件系统访问",
        "由用户在本地选择文件或目录",
    ],
]

# 对照《软件需求规格说明书》《软件概要设计说明书》《软件详细设计说明书》
# 章节编号与 SRS §6.1 追溯矩阵、LLD §2672 追溯矩阵保持一致。
TRACE_ROWS = [
    [
        "1",
        "SRA-001",
        "FR-001; SR-001",
        "3.1 登录模块",
        "3.2.4 认证流程",
        "UserManager.cxx; qLoginDialog.cxx",
        "T3,T4,T7",
    ],
    [
        "2",
        "SRA-002",
        "FR-003; SR-010",
        "7.1 出错信息; 8.2 安全与隐私",
        "3.2.2 安全模型; 3.2.4 认证流程",
        "UserManager.cxx (authenticate/lock)",
        "T13,T14",
    ],
    [
        "3",
        "SRA-003",
        "FR-002; SR-005",
        "8.2 安全与隐私设计",
        "3.2.2 安全模型; 3.2.6 与主窗口集成",
        "qUserManagementDialog.cxx; qRadianceAppMainWindow.cxx",
        "T9,T10",
    ],
    [
        "4",
        "SRA-004",
        "FR-004; SR-002",
        "3.1 设置模块/用户设置",
        "3.2.3 密码存储; 3.2.5 界面类",
        "qChangePasswordDialog.cxx; UserManager.cxx",
        "T8,T28",
    ],
    [
        "5",
        "SRA-005",
        "FR-005; DR-002",
        "8.1.2 用户数据维护",
        "3.2.5 界面类; 3.2.6 与主窗口集成",
        "qUserManagementDialog.cxx; UserManager.cxx",
        "T10,T14,T27",
    ],
    [
        "6",
        "SRA-006",
        "SR-001（启动约束）",
        "2.1 系统架构设计",
        "2.4 启动阶段",
        "Main.cxx (QLockFile 单实例)",
        "T5",
    ],
    [
        "7",
        "SRA-007",
        "FR-070; SR-004",
        "5.2 内部接口（授权服务）",
        "3.3 授权系统",
        "LicenseManager.cxx; qLicenseCheckDialog.cxx",
        "T6",
    ],
    [
        "8",
        "SRA-008",
        "FR-058; FR-100; SR-006",
        "11.1 系统退出",
        "8 系统退出; 3.2.6 与主窗口集成",
        "qRadianceAppMainWindow.cxx (logout)",
        "T1,T15",
    ],
    [
        "9",
        "SRA-009",
        "FR-060; DR-004",
        "8.1.1 日志维护",
        "3.6 日志; 7.1 日志点位清单",
        "AppLogger.cxx; Main.cxx",
        "T2",
    ],
    [
        "10",
        "SRA-010",
        "FR-020; DR-001; DR-008",
        "2.2 影像导入流程; 5.3 外部接口",
        "3.4 影像导入校验; §2.6 数据主路径",
        "Slicer DICOM 模块; VisionMagicConfig.py",
        "T17,T18,T26",
    ],
    [
        "11",
        "SRA-011",
        "FR-021; RR-002",
        "4.1.3 数据一致性与恢复",
        "3.4.3 失败语义; 3.4.4 与 MRML 生命周期",
        "VolumeImportValidator.cxx",
        "T19",
    ],
    [
        "12",
        "SRA-012",
        "RR-001（删除二次确认）",
        "7.1 出错信息与补救措施",
        "6.1 错误分类",
        "Slicer 数据/场景模块",
        "T20",
    ],
    [
        "13",
        "SRA-013",
        "FR-050; DR-007",
        "2.2/4.1.2 持久化设计",
        "§2.6 医生标注与保存",
        "qRadianceAppMainWindow.cxx (保存)",
        "T21",
    ],
    [
        "14",
        "SRA-014",
        "FR-001; FR-005",
        "3.1 登录模块（格式校验）",
        "3.2.3 密码存储; 3.2.4 认证流程",
        "UserManager.cxx (validateUsername/limits)",
        "T11,T12",
    ],
    [
        "15",
        "SRA-015",
        "SR-003; SR-007; SR-009",
        "8.2/8.3 安全与隐私; 部署配置",
        "7.2 配置项清单; 7.3 非功能落实; 9.2 运行环境",
        "产品说明书",
        "T22,T23,T24,T25,T29,T32,T33",
    ],
    [
        "16",
        "SRA-016",
        "—（SBOM/RDMP 文档要求）",
        "8.1.4 版本信息维护",
        "附录 B 实现映射表",
        "现成软件清单文档",
        "T30,T31",
    ],
    [
        "17",
        "SRA-017",
        "SR-002; FR-001",
        "3.1 登录模块（密文输入）",
        "3.2.5 界面类（qLoginDialog）",
        "qLoginDialog.cxx (Password echo)",
        "T16",
    ],
    [
        "18",
        "SRA-018",
        "SR-008（运行环境兼容）",
        "3.2 运行环境",
        "9.2 运行环境",
        "—（说明书+杀毒软件共存验证）",
        "T34",
    ],
]


def _replace_in_text(text: str) -> str:
    rules = [
        ("**有限公司", "（注册人名称）有限公司"),
        ("手术计划软件", SOFTWARE_NAME),
        ("***软件", SOFTWARE_NAME),
        ("**软件", SOFTWARE_NAME),
        ("***影像", "DICOM 影像"),
        ("****", "设备"),
        ("**使用者", "使用者"),
        ("脑部磁共振影像数据处理软件", SOFTWARE_NAME),
        ("V1.0.0.5", VERSION),
        ("V1.0.0.0", VERSION),
        (
            "申报产品通过局域网的DICOM3.0接口的DICOM3.0协议获取到***影像。",
            "申报产品以单机本地运行为主，用户通过本地文件系统或 DICOM 模块导入 DICOM 影像数据，"
            "在软件内完成浏览、三维重建及相关处理，正常运行不依赖外部网络连接。",
        ),
        (
            "1.1.3 手术计划软件运行环境",
            "1.1.3 软件运行环境",
        ),
        (
            "***软件使用加密狗、用户名和密码进行访问限制。",
            f"{SOFTWARE_NAME}使用用户名和密码进行访问控制；"
            "可按产品部署策略启用授权/许可校验（如加密狗）。",
        ),
        (
            "管理员权限：访问配置程序、数据库、程序维护",
            "管理员权限：用户管理（添加/锁定/解锁/重置密码）、系统维护",
        ),
        (
            "普通用户访问权限：可以操作**软件、更改用户名和密码、切换语言。",
            "普通用户权限：操作影像浏览与三维重建功能、修改自身密码、退出登录。",
        ),
        (
            "配置端管理员账号",
            "管理员账号",
        ),
        (
            "通过输入正确的用户名和密码成功登录后，才可以使用手术计划软件用户端。",
            f"通过输入正确的用户名和密码成功登录后，才可以使用{SOFTWARE_NAME}。",
        ),
        (
            "通过管理员密码方可进入手术计划软件配置端。",
            "管理员通过【用户管理】进行账户与权限管理。",
        ),
        (
            "修改用户名和密码，需再次输入登录密码确认后方可修改。",
            "修改密码需输入原密码确认；管理员可重置用户密码。",
        ),
        (
            "详见《网络安全测试记录》",
            f"详见《{SOFTWARE_NAME}-网络安全测试用例-修订.docx》",
        ),
        (
            "在进行软件测试和网络安全测试过程中，均在此安全软件运行的环境下进行测试工作，其兼容性满足要求。"
            "本产品在进行数据传输过程中，使用了标准协议DICOM、TCP/IP，数据存储格式为标准的dcm，jpeg，json和mgjson，"
            "并提供了真实性声明，详见《网络安全测试计划》、《网络安全测试报告》、《标准传输协议及存储格式声明》。",
            "在进行软件测试和网络安全测试过程中，均在火绒等安全软件运行的环境下进行，兼容性满足要求。"
            "本产品以本地 DICOM/影像文件导入为主，数据存储格式包括 dcm、nrrd、nii、mrml、mrb 等，"
            f"详见《网络安全测试计划》、《网络安全测试报告》、《{SOFTWARE_NAME}-网络安全测试用例-修订.docx》。",
        ),
    ]
    for old, new in rules:
        text = text.replace(old, new)
    return text


def _apply_text_replacements(doc: Document) -> None:
    for para in doc.paragraphs:
        if para.text.strip():
            new = _replace_in_text(para.text)
            if new != para.text:
                para.text = new
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    cell.text = _replace_in_text(cell.text)


def _set_cell(cell, text: str) -> None:
    cell.text = text


def _rewrite_table_rows(table, header_rows: int, data_rows: list[list[str]]) -> None:
    while len(table.rows) > header_rows:
        table._tbl.remove(table.rows[-1]._tr)
    for values in data_rows:
        row = table.add_row()
        for i, val in enumerate(values):
            _set_cell(row.cells[i], val)


def revise_research_report(src: Path, out: Path) -> None:
    doc = Document(str(src))

    # 关键段落定点替换
    for para in doc.paragraphs:
        t = para.text.strip()
        if t == "客户端软件名称：":
            para.text = f"客户端软件名称：{SOFTWARE_NAME}（{SOFTWARE_EN}）"
        elif t == "客户端软件发布版本：V1.0":
            para.text = f"客户端软件发布版本：{VERSION}"
        elif t == "规格型号：":
            para.text = f"规格型号：{MODEL}"
        elif t == "注册人/生产企业名称：":
            para.text = "注册人/生产企业名称：（请填写）"
        elif t == "注册人/生产企业住所：":
            para.text = "注册人/生产企业住所：（请填写）"

    _apply_text_replacements(doc)

    # TABLE 0 运行环境 — 保持，仅确保 OS 行正确
    # TABLE 1 通讯数据
    if len(doc.tables) > 1:
        t1 = doc.tables[1]
        _set_cell(t1.rows[0].cells[0], "数据类型")
        _rewrite_table_rows(t1, 1, COMM_DATA_ROWS)

    # TABLE 2 电子接口
    if len(doc.tables) > 2:
        t2 = doc.tables[2]
        _rewrite_table_rows(t2, 1, INTERFACE_ROWS)

    # TABLE 3 网络安全能力
    if len(doc.tables) > 3:
        t3 = doc.tables[3]
        _rewrite_table_rows(t3, 1, CAPABILITY_ROWS)

    # TABLE 4 安全软件 — 去掉空加密狗行，保留火绒
    if len(doc.tables) > 4:
        t4 = doc.tables[4]
        _set_cell(t4.rows[1].cells[0], "火绒安全软件")
        _set_cell(t4.rows[1].cells[1], "（现场填写版本）")
        _set_cell(t4.rows[1].cells[2], "北京火绒网络科技有限公司")
        _set_cell(t4.rows[1].cells[3], "Windows 10/11，x64")

    doc.save(str(out))
    print(f"研究报告: {out}")


def revise_test_report(src: Path, out: Path) -> None:
    doc = Document(str(src))
    _apply_text_replacements(doc)

    for para in doc.paragraphs:
        t = para.text.strip()
        if t == "产品名称：":
            para.text = f"产品名称：{SOFTWARE_NAME}"
        elif t == "产品型号：":
            para.text = f"产品型号：{MODEL}"
        elif t.startswith("测试版本："):
            para.text = f"测试版本：{VERSION}"

    if doc.tables:
        # TABLE 0 测试时间
        t0 = doc.tables[0]
        _set_cell(t0.rows[1].cells[0], MODEL)

        # TABLE 1 测试环境
        if len(doc.tables) > 1:
            t1 = doc.tables[1]
            env = {
                1: "Windows 10/11 64 位",
                2: "主频 2.2GHz 及以上",
                3: "16GB 及以上",
                4: "500GB 及以上（剩余安装空间 10GB 以上）",
                5: "标准输入设备",
                6: "1920×1080 及以上",
                7: "显存 8GB 及以上（NVIDIA RTX 3050 或同等级别）",
            }
            for ri, val in env.items():
                if ri < len(t1.rows):
                    _set_cell(t1.rows[ri].cells[2], val)

        # TABLE 2 测试工具 — 去掉 dcm4che，改为测试数据
        if len(doc.tables) > 2:
            t2 = doc.tables[2]
            _set_cell(t2.rows[1].cells[1], "合规 DICOM 测试数据集")
            _set_cell(t2.rows[1].cells[2], "—")
            _set_cell(t2.rows[1].cells[3], "本地 DICOM 导入与浏览测试")
            _set_cell(t2.rows[2].cells[1], "测试工作站")
            _set_cell(t2.rows[2].cells[2], "Windows 10/11")
            _set_cell(t2.rows[2].cells[3], "软件功能与网络安全测试")

        # TABLE 3 安全软件 — 移除加密狗行
        if len(doc.tables) > 3:
            t3 = doc.tables[3]
            while len(t3.rows) > 2:
                t3._tbl.remove(t3.rows[-1]._tr)
            _set_cell(t3.rows[1].cells[1], "火绒安全软件")
            _set_cell(t3.rows[1].cells[2], "（现场填写版本）")
            _set_cell(t3.rows[1].cells[3], "北京火绒网络科技有限公司")
            _set_cell(t3.rows[1].cells[4], "Windows 10/11")

    doc.save(str(out))
    print(f"测试报告: {out}")


def revise_traceability_doc(src: Path, out_docx: Path) -> None:
    """将 .doc 追溯表另存为 .docx 并更新内容。"""
    try:
        import win32com.client  # type: ignore

        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(str(src))

        # 标题段落
        if doc.Paragraphs.Count >= 3:
            doc.Paragraphs(1).Range.Text = "软件网络安全可追溯性分析记录\r"
            doc.Paragraphs(3).Range.Text = (
                f"软件名称：{SOFTWARE_NAME}（{SOFTWARE_EN}）"
                f"                                                         软件版本号：{VERSION}\r"
            )

        table = doc.Tables(1)
        headers = [
            "序号",
            "风险管理编号",
            "软件需求编号",
            "概要设计说明书编号",
            "详细设计说明书编号",
            "源代码\n（追溯源代码和系统测试用例的对应关系）",
            "网络安全测试编号",
        ]
        for ci, h in enumerate(headers, start=1):
            table.Cell(1, ci).Range.Text = h

        # 删除多余行
        while table.Rows.Count > len(TRACE_ROWS) + 1:
            table.Rows(len(TRACE_ROWS) + 2).Delete()

        # 确保行数足够
        while table.Rows.Count < len(TRACE_ROWS) + 1:
            table.Rows.Add()

        for ri, row in enumerate(TRACE_ROWS, start=2):
            for ci, val in enumerate(row, start=1):
                table.Cell(ri, ci).Range.Text = val

        # 更新脚注段落（若存在）
        for i in range(1, doc.Paragraphs.Count + 1):
            t = doc.Paragraphs(i).Range.Text
            if "记录时间" in t:
                doc.Paragraphs(i).Range.Text = (
                    "记录时间：                   （现场填写）         "
                    "追溯人：                       审核人：\r"
                )
            elif t.startswith("注："):
                doc.Paragraphs(i).Range.Text = (
                    "注：RM 编码来自风险管理资料；软件需求编号来自《软件需求规格说明书》"
                    "（含 FR/DR/SR/RR）；概要/详细设计章节编号分别来自《软件概要设计说明书》"
                    "《软件详细设计说明书》；网络安全测试编号对应《网络安全测试用例》T 系列用例。"
                    "每项风险可与需求、设计、源代码、测试一对一或多对一对应。\r"
                )

        doc.SaveAs2(str(out_docx), FileFormat=16)  # wdFormatDocumentDefault = 16 (.docx)
        doc.Close(False)
        word.Quit()
        print(f"追溯表: {out_docx}")
    except Exception as exc:
        print(f"win32com 修订追溯表失败: {exc}；改生成新 docx")
        _create_traceability_docx(out_docx)


def _create_traceability_docx(out: Path) -> None:
    doc = Document()
    doc.add_paragraph("软件网络安全可追溯性分析记录")
    doc.add_paragraph("编号：SKAR-JL-241-000")
    doc.add_paragraph(
        f"软件名称：{SOFTWARE_NAME}（{SOFTWARE_EN}）"
        f"                                                         软件版本号：{VERSION}"
    )
    headers = TRACE_ROWS[0] if False else [
        "序号",
        "风险管理编号",
        "软件需求编号",
        "概要设计说明书编号",
        "详细设计说明书编号",
        "源代码（追溯源代码和系统测试用例的对应关系）",
        "网络安全测试编号",
    ]
    table = doc.add_table(rows=1 + len(TRACE_ROWS), cols=7)
    table.style = "Table Grid"
    for j, h in enumerate(headers):
        table.rows[0].cells[j].text = h
    for i, row in enumerate(TRACE_ROWS, start=1):
        for j, val in enumerate(row):
            table.rows[i].cells[j].text = val
    doc.add_paragraph(
        "记录时间：                   （现场填写）         "
        "追溯人：                       审核人："
    )
    doc.add_paragraph(
        "注：RM 编码来自风险管理资料；软件需求编号来自《软件需求规格说明书》"
        "（含 FR/DR/SR/RR）；概要/详细设计章节编号分别来自《软件概要设计说明书》"
        "《软件详细设计说明书》；网络安全测试编号对应《网络安全测试用例》T 系列用例。"
        "每项风险可与需求、设计、源代码、测试一对一或多对一对应。"
    )
    doc.save(str(out))
    print(f"追溯表(新建): {out}")


def _backup(path: Path) -> None:
    bak = path.with_name(path.stem + "-原稿备份" + path.suffix)
    if path.exists() and not bak.exists():
        shutil.copy2(path, bak)
        print(f"备份: {bak}")


def main() -> None:
    research_src = BASE / "CH3.5.5.11-网络安全研究报告.docx"
    report_src = BASE / "CH3.5.5.11.2-网络安全测试报告.docx"
    trace_src = BASE / "ch3.5.5.11网络安全可追溯性分析汇总表1110.doc"

    research_out = BASE / "CH3.5.5.11-网络安全研究报告-修订.docx"
    report_out = BASE / "CH3.5.5.11.2-网络安全测试报告-修订.docx"
    trace_out = BASE / "ch3.5.5.11网络安全可追溯性分析汇总表-修订.docx"

    for p in (research_src, report_src, trace_src):
        if not p.exists():
            raise FileNotFoundError(p)
        _backup(p)

    revise_research_report(research_src, research_out)
    revise_test_report(report_src, report_out)
    revise_traceability_doc(trace_src, trace_out)

    # 尝试覆盖原文件
    for src, out in (
        (research_src, research_out),
        (report_src, report_out),
    ):
        try:
            shutil.copy2(out, src)
            print(f"已更新原文件: {src.name}")
        except PermissionError:
            print(f"原文件被占用，请手动替换: {src.name}")

    print("完成。")


if __name__ == "__main__":
    main()
