# -*- coding: utf-8 -*-
from pathlib import Path

BASE = Path(r"C:\Users\lie76\Desktop\三维重建软件文档\过审材料\软件资料需求2026.06.23")
OUT = Path(__file__).resolve().parent / "_ch_docs_dump.txt"

files = [
    BASE / "CH3.5.5.11-网络安全研究报告.docx",
    BASE / "CH3.5.5.11.2-网络安全测试报告.docx",
    BASE / "ch3.5.5.11网络安全可追溯性分析汇总表1110.doc",
]

lines = []

def dump_docx(path: Path) -> None:
    from docx import Document
    doc = Document(str(path))
    lines.append(f"\n{'='*80}\nFILE: {path.name}\n{'='*80}")
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            lines.append(f"P{i}: {para.text}")
    for ti, table in enumerate(doc.tables):
        lines.append(f"\n--- TABLE {ti}: {len(table.rows)}x{len(table.columns)} ---")
        for ri, row in enumerate(table.rows):
            cells = [c.text.replace("\n", " | ") for c in row.cells]
            lines.append(f"R{ri}: " + " || ".join(cells))


def dump_doc(path: Path) -> None:
    try:
        import win32com.client  # type: ignore
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(str(path))
        lines.append(f"\n{'='*80}\nFILE: {path.name}\n{'='*80}")
        for i in range(1, doc.Paragraphs.Count + 1):
            t = doc.Paragraphs(i).Range.Text.strip()
            if t and t != "\r":
                lines.append(f"P{i}: {t}")
        for ti in range(1, doc.Tables.Count + 1):
            table = doc.Tables(ti)
            lines.append(f"\n--- TABLE {ti-1}: {table.Rows.Count}x{table.Columns.Count} ---")
            for ri in range(1, table.Rows.Count + 1):
                cells = []
                for ci in range(1, table.Columns.Count + 1):
                    cells.append(table.Cell(ri, ci).Range.Text.replace("\r\x07", "").replace("\n", " | "))
                lines.append(f"R{ri-1}: " + " || ".join(cells))
        doc.Close(False)
        word.Quit()
        return
    except Exception as e:
        lines.append(f"win32com failed: {e}")

    # fallback: antiword-like via textract not available; try olefile / raw
    try:
        from docx import Document
        dump_docx(path)
    except Exception as e2:
        lines.append(f"docx fallback failed: {e2}")


for f in files:
    if not f.exists():
        lines.append(f"MISSING: {f}")
        continue
    if f.suffix.lower() == ".docx":
        dump_docx(f)
    else:
        dump_doc(f)

OUT.write_text("\n".join(lines), encoding="utf-8")
print(OUT)
