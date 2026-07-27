# -*- coding: utf-8 -*-
from pathlib import Path
from docx import Document

BASE = Path(r"C:\Users\lie76\Desktop\三维重建软件文档\过审材料\软件资料需求2026.06.23")
files = [
    BASE / "CH3.5.5.11-网络安全研究报告-修订.docx",
    BASE / "CH3.5.5.11.2-网络安全测试报告-修订.docx",
    BASE / "ch3.5.5.11网络安全可追溯性分析汇总表-修订.docx",
]
out = Path(__file__).resolve().parent / "_ch_revised_dump.txt"
lines = []
for f in files:
    doc = Document(str(f))
    lines.append(f"\n=== {f.name} ===")
    for p in doc.paragraphs[:40]:
        if p.text.strip():
            lines.append(p.text[:200])
    for ti, t in enumerate(doc.tables):
        lines.append(f"-- T{ti} {len(t.rows)}x{len(t.columns)} --")
        for ri in range(min(5, len(t.rows))):
            lines.append(" | ".join(c.text[:80].replace("\n"," ") for c in t.rows[ri].cells))
        if len(t.rows) > 5:
            lines.append(f"... +{len(t.rows)-5} rows")
out.write_text("\n".join(lines), encoding="utf-8")
print(out)
